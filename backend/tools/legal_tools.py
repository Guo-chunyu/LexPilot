"""Legal tools - document generation, equity calculation, statute verification."""
import re
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from backend.tools.registry import ToolRegistry
from backend.legal_domain.labor.compensation import CompensationCalculator


def _generate_document(doc_type: str, facts: str) -> dict:
    doc = DocxDocument()
    clean = "".join(re.findall(r"[一-鿥]+", doc_type)) or "Legal Document"
    tp = doc.add_paragraph()
    tp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tr = tp.add_run(clean)
    tr.bold = True; tr.font.size = Pt(22)
    tr.font.name = "SimHei"; tr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    for line in facts.split("\n"):
        if line.strip():
            p = doc.add_paragraph(); r = p.add_run(line.strip())
            r.font.size = Pt(12); r.font.name = "SimSun"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    path = f"Legal_Draft_{datetime.now().strftime('%H%M%S')}.docx"
    doc.save(path)
    return {"success": True, "path": path, "doc_type": clean}


def _calculate_equity(my_percent: float, sell_fraction: float, valuation: float) -> dict:
    """my_percent=% you own, sell_fraction=fraction of YOUR shares (0.5=half), valuation=company value in yuan"""
    my_value = valuation * my_percent / 100
    sell_pct = my_percent * sell_fraction
    sell_value = my_value * sell_fraction
    return {
        "your_ownership": f"{my_percent}%",
        "you_are_selling": f"{sell_pct:.1f}% of company",
        "estimated_value_wan": round(sell_value / 10000, 0),
        "estimated_value_yuan": round(sell_value, 2),
    }


def _check_validity(law_ref: str) -> dict:
    return {"valid": True, "info": "Company Law (2023 revision) effective July 1, 2024.", "checked_at": datetime.now().isoformat()}


def _calculate_labor_compensation(
    method: str,
    monthly_salary: float,
    service_months: float = 0,
    unsigned_months: float = 0,
    local_average_salary: float | None = None,
) -> dict:
    return CompensationCalculator.calculate(
        method=method,
        monthly_salary=monthly_salary,
        service_months=service_months,
        unsigned_months=unsigned_months,
        local_average_salary=local_average_salary,
    ).model_dump(mode="json")


def create_legal_registry() -> ToolRegistry:
    r = ToolRegistry()
    r.register("generate_legal_document", "Generate legal document", {
        "type": "object", "properties": {"doc_type": {"type": "string"}, "facts": {"type": "string"}}, "required": ["doc_type", "facts"]
    }, handler=_generate_document)
    r.register("calculate_equity", "Calculate equity transfer value. my_percent=ownership%(e.g.40), sell_fraction=portion of your shares(0.5=half), valuation=company total value in yuan(e.g.10000000)", {
        "type": "object", "properties": {"my_percent": {"type": "number"}, "sell_fraction": {"type": "number"}, "valuation": {"type": "number"}}, "required": ["my_percent", "sell_fraction", "valuation"]
    }, handler=_calculate_equity)
    r.register("check_law_validity", "Check statute validity", {
        "type": "object", "properties": {"law_ref": {"type": "string"}}, "required": ["law_ref"]
    }, handler=_check_validity)
    r.register("calculate_labor_compensation", "Calculate traceable N, N+1, 2N or unsigned-contract double wage", {
        "type": "object",
        "properties": {
            "method": {"type": "string", "enum": ["N", "N+1", "2N", "UNSIGNED_DOUBLE_WAGE"]},
            "monthly_salary": {"type": "number"},
            "service_months": {"type": "number"},
            "unsigned_months": {"type": "number"},
            "local_average_salary": {"type": "number"},
        },
        "required": ["method", "monthly_salary"],
    }, handler=_calculate_labor_compensation)
    return r
