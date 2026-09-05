"""Secure local ingestion for user-provided labor-dispute evidence files."""

from __future__ import annotations

import csv
import hashlib
import io
import mimetypes
import re
from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path
from zipfile import BadZipFile, ZipFile

import yaml

from backend.legal_domain.labor.evidence_gap import detect_evidence_gaps
from backend.legal_domain.labor.facts import extract_labor_facts
from backend.legal_rl.state import (
    CaseState,
    FileExtractionStatus,
    UploadedEvidenceFile,
)


MAX_FILES_PER_MESSAGE = 10
MAX_FILE_BYTES = 15 * 1024 * 1024
MAX_EXTRACTED_CHARS = 50_000
MAX_ZIP_UNCOMPRESSED_BYTES = 60 * 1024 * 1024
ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".txt",
    ".md",
    ".csv",
    ".xlsx",
    ".xls",
}
STREAMLIT_FILE_TYPES = sorted(extension.removeprefix(".") for extension in ALLOWED_EXTENSIONS)
CURRENT_EVIDENCE_PARSER_VERSION = 3


@dataclass(frozen=True)
class UploadPayload:
    name: str
    data: bytes
    media_type: str = "application/octet-stream"


@dataclass
class IngestionResult:
    files: list[UploadedEvidenceFile] = field(default_factory=list)
    duplicate_names: list[str] = field(default_factory=list)

    @property
    def extracted_count(self) -> int:
        return sum(
            item.extraction_status == FileExtractionStatus.TEXT_EXTRACTED
            for item in self.files
        )

    def summary_markdown(self) -> str:
        lines = [f"已接收并登记 {len(self.files)} 份材料："]
        for item in self.files:
            labels = "、".join(item.evidence_names) or "待人工分类"
            status = {
                FileExtractionStatus.TEXT_EXTRACTED: "已提取文字",
                FileExtractionStatus.METADATA_ONLY: "已留存，暂无可提取文字",
                FileExtractionStatus.PARSER_UNAVAILABLE: "已留存，本机暂无对应解析器",
                FileExtractionStatus.FAILED: "已留存，解析失败",
            }[item.extraction_status]
            details = ""
            if item.extracted_facts:
                details = "；从正文识别：" + "、".join(item.extracted_facts)
            elif item.extraction_status == FileExtractionStatus.TEXT_EXTRACTED:
                details = "；已读取正文，暂未识别到可直接入库的关键事实"
            lines.append(f"- `{item.original_name}`：{status}；材料类型：{labels}{details}")
        if self.duplicate_names:
            lines.append("- 重复材料未重复登记：" + "、".join(self.duplicate_names))
        return "\n".join(lines)


class EvidenceUploadError(ValueError):
    """Raised when an upload fails allow-list or safety validation."""


def ingest_evidence_files(
    state: CaseState,
    uploads: list[UploadPayload],
    storage_root: str | Path,
) -> IngestionResult:
    """Validate, store, extract and attach uploaded evidence to one CaseState."""

    if not uploads:
        return IngestionResult()
    if len(uploads) > MAX_FILES_PER_MESSAGE:
        raise EvidenceUploadError(f"单次最多上传 {MAX_FILES_PER_MESSAGE} 个文件。")

    root = Path(storage_root).resolve()
    safe_case_id = re.sub(r"[^A-Za-z0-9_-]", "_", state.case_id)[:80] or "case"
    case_dir = (root / safe_case_id).resolve()
    if case_dir != root and root not in case_dir.parents:
        raise EvidenceUploadError("案件材料存储路径无效。")
    case_dir.mkdir(parents=True, exist_ok=True)

    pending_requests = list(state.pending_evidence_requests)
    existing_by_hash = {item.sha256: item for item in state.uploaded_files}
    existing_text_fingerprints = {
        fingerprint
        for item in state.uploaded_files
        if (fingerprint := _text_fingerprint(item.extension, item.text_preview))
    }
    result = IngestionResult()

    for index, upload in enumerate(uploads):
        original_name = _safe_original_name(upload.name)
        extension = Path(original_name).suffix.lower()
        _validate_upload(original_name, extension, upload.data)
        digest = hashlib.sha256(upload.data).hexdigest()
        if digest in existing_by_hash:
            result.duplicate_names.append(original_name)
            continue

        text, status, page_count, notes = _extract_text(extension, upload.data)
        from backend.legal_domain.consultation.profiles import route_case, classify_material
        route_case(original_name + "\n" + text[:600], state)
        text_fingerprint = _text_fingerprint(extension, text)
        if text_fingerprint and text_fingerprint in existing_text_fingerprints:
            result.duplicate_names.append(original_name)
            continue

        stored_name = f"{digest[:24]}{extension}"
        stored_path = case_dir / stored_name
        if not stored_path.exists():
            stored_path.write_bytes(upload.data)

        evidence_names = _classify_evidence(
            original_name,
            text,
            pending_requests,
            file_index=index,
            file_count=len(uploads),
        )
        if state.case_type != "labor_dispute":
            evidence_names = classify_material(original_name, text)
        media_type = upload.media_type or mimetypes.guess_type(original_name)[0] or "application/octet-stream"
        record = UploadedEvidenceFile(
            file_id=f"file_{digest[:16]}",
            original_name=original_name,
            stored_name=stored_name,
            stored_path=str(stored_path),
            media_type=media_type,
            extension=extension,
            size_bytes=len(upload.data),
            sha256=digest,
            evidence_names=evidence_names,
            extraction_status=status,
            text_preview=_preview(text),
            page_count=page_count,
            notes=notes,
        )
        state.uploaded_files.append(record)
        existing_by_hash[digest] = record
        if text_fingerprint:
            existing_text_fingerprints.add(text_fingerprint)
        result.files.append(record)

        for evidence_name in evidence_names:
            state.add_evidence(
                evidence_name,
                source="uploaded_file",
                notes=f"{original_name} · sha256:{digest[:16]}",
            )
        if text:
            # Raw uploaded materials remain local by default. Semantic enrichment is
            # reserved for the user's chat text unless a future explicit consent flow
            # is added for remote document processing.
            _merge_extracted_text(state, record, text)

    state.evidence_parser_version = CURRENT_EVIDENCE_PARSER_VERSION
    _refresh_domain_evidence(state)
    return result


def refresh_stored_evidence(
    state: CaseState,
    storage_root: str | Path,
) -> bool:
    """Re-index stored files once when an older live session uses a new parser."""

    if state.evidence_parser_version >= CURRENT_EVIDENCE_PARSER_VERSION:
        return False
    if not state.uploaded_files:
        state.evidence_parser_version = CURRENT_EVIDENCE_PARSER_VERSION
        return False

    root = Path(storage_root).resolve()
    case_dir = (root / re.sub(r"[^A-Za-z0-9_-]", "_", state.case_id)[:80]).resolve()
    previews = [item.text_preview for item in state.uploaded_files if item.text_preview]
    stale_fact_ids: set[str] = set()
    retained_provenance = []
    filenames = {item.original_name for item in state.uploaded_files}
    for item in state.fact_provenance:
        old_unlabelled_file_trace = (
            item.source_type == "user_message"
            and not item.source_ref
            and item.quote
            and any(item.quote in preview for preview in previews)
        )
        if (item.source_type == "uploaded_file" and item.source_ref in filenames) or old_unlabelled_file_trace:
            stale_fact_ids.add(item.fact_id)
        else:
            retained_provenance.append(item)
    state.fact_provenance = retained_provenance
    retained_fact_ids = {item.fact_id for item in retained_provenance if item.accepted}
    for fact_id in stale_fact_ids - retained_fact_ids:
        state.facts.pop(fact_id, None)

    refreshed: list[tuple[UploadedEvidenceFile, str]] = []
    valid_uploaded_evidence: set[str] = set()
    for record in state.uploaded_files:
        stored_path = Path(record.stored_path) if record.stored_path else case_dir / record.stored_name
        if not stored_path.is_file():
            continue
        data = stored_path.read_bytes()
        text, status, page_count, notes = _extract_text(record.extension, data)
        evidence_names = _classify_evidence(
            record.original_name,
            text,
            [],
            file_index=0,
            file_count=1,
        )
        if state.case_type != "labor_dispute":
            from backend.legal_domain.consultation.profiles import classify_material
            evidence_names = classify_material(record.original_name, text)
        record.stored_path = str(stored_path)
        record.evidence_names = evidence_names
        record.extraction_status = status
        record.text_preview = _preview(text)
        record.page_count = page_count
        record.notes = notes
        record.extracted_facts = []
        valid_uploaded_evidence.update(evidence_names)
        refreshed.append((record, text))

    state.evidence = [
        item for item in state.evidence
        if item.source != "uploaded_file" or item.name in valid_uploaded_evidence
    ]
    for record, text in refreshed:
        for evidence_name in record.evidence_names:
            state.add_evidence(
                evidence_name,
                source="uploaded_file",
                notes=f"{record.original_name} · sha256:{record.sha256[:16]}",
            )
        if text:
            _merge_extracted_text(state, record, text)

    state.evidence_parser_version = CURRENT_EVIDENCE_PARSER_VERSION
    state.judge_result = None
    state.verification_result = None
    state.final_report = {}
    _refresh_domain_evidence(state)
    return bool(refreshed)


def _merge_extracted_text(state: CaseState, record: UploadedEvidenceFile, text: str) -> None:
    if state.case_type != "labor_dispute":
        from backend.legal_domain.consultation.intake import ingest_text, LABELS
        before = dict(state.facts)
        ingest_text(text, state, source_type="uploaded_file", source_ref=record.original_name, contextual=False)
        record.extracted_facts = [f"{LABELS.get(key, key)}：{value}" for key, value in state.facts.items() if before.get(key) != value]
        return
    provenance_start = len(state.fact_provenance)
    if "劳动合同" in record.evidence_names:
        state.apply_facts({"has_written_contract": True})
        state.add_fact_provenance(
            "has_written_contract",
            source_type="uploaded_file",
            source_ref=record.original_name,
            quote=_contract_identity_quote(text),
            extraction_method="document_type",
        )
    extract_labor_facts(
        text,
        state,
        enable_semantic=False,
        source_type="uploaded_file",
        source_ref=record.original_name,
        recognize_evidence_mentions=False,
        interpret_pending_answer=False,
        document_evidence_names=tuple(record.evidence_names),
    )
    fact_ids = list(dict.fromkeys(
        item.fact_id
        for item in state.fact_provenance[provenance_start:]
        if item.source_type == "uploaded_file" and item.source_ref == record.original_name
    ))
    record.extracted_facts = [
        _format_extracted_fact(fact_id, state.facts.get(fact_id))
        for fact_id in fact_ids
        if state.facts.get(fact_id) is not None
    ]


def _refresh_domain_evidence(state: CaseState) -> None:
    if state.case_type == "labor_dispute":
        detect_evidence_gaps(state)
    else:
        from backend.legal_domain.consultation.intake import refresh_evidence
        refresh_evidence(state)


def _safe_original_name(name: str) -> str:
    normalized = str(name or "").replace("\\", "/")
    basename = normalized.rsplit("/", 1)[-1].strip().replace("\x00", "")
    if not basename or basename in {".", ".."}:
        raise EvidenceUploadError("文件名无效。")
    return basename[:180]


def _validate_upload(name: str, extension: str, data: bytes) -> None:
    if extension not in ALLOWED_EXTENSIONS:
        allowed = "、".join(sorted(ALLOWED_EXTENSIONS))
        raise EvidenceUploadError(f"不支持 `{name}` 的文件类型。允许：{allowed}")
    if not data:
        raise EvidenceUploadError(f"`{name}` 是空文件。")
    if len(data) > MAX_FILE_BYTES:
        raise EvidenceUploadError(f"`{name}` 超过单文件 15 MB 限制。")

    if extension == ".pdf" and not data.startswith(b"%PDF"):
        raise EvidenceUploadError(f"`{name}` 不是有效的 PDF 文件。")
    if extension == ".png" and not data.startswith(b"\x89PNG\r\n\x1a\n"):
        raise EvidenceUploadError(f"`{name}` 不是有效的 PNG 文件。")
    if extension in {".jpg", ".jpeg"} and not data.startswith(b"\xff\xd8\xff"):
        raise EvidenceUploadError(f"`{name}` 不是有效的 JPEG 文件。")
    if extension == ".webp" and not (data.startswith(b"RIFF") and data[8:12] == b"WEBP"):
        raise EvidenceUploadError(f"`{name}` 不是有效的 WebP 文件。")
    if extension in {".doc", ".xls"} and not data.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
        raise EvidenceUploadError(f"`{name}` 不是有效的旧版 Office 文件。")
    if extension in {".docx", ".xlsx"}:
        _validate_office_zip(name, extension, data)


def _validate_office_zip(name: str, extension: str, data: bytes) -> None:
    try:
        with ZipFile(io.BytesIO(data)) as archive:
            infos = archive.infolist()
            total_size = sum(info.file_size for info in infos)
            if len(infos) > 3000 or total_size > MAX_ZIP_UNCOMPRESSED_BYTES:
                raise EvidenceUploadError(f"`{name}` 解压后体积或文件数量异常。")
            members = {info.filename for info in infos}
            required = "word/document.xml" if extension == ".docx" else "xl/workbook.xml"
            if "[Content_Types].xml" not in members or required not in members:
                raise EvidenceUploadError(f"`{name}` 不是有效的 {extension.upper()} 文件。")
    except BadZipFile as exc:
        raise EvidenceUploadError(f"`{name}` 不是有效的 Office 文件。") from exc


def _extract_text(
    extension: str,
    data: bytes,
) -> tuple[str, FileExtractionStatus, int | None, str]:
    try:
        if extension == ".pdf":
            return _extract_pdf(data)
        if extension == ".docx":
            return _extract_docx(data)
        if extension == ".xlsx":
            return _extract_xlsx(data)
        if extension == ".csv":
            text = _decode_text(data)
            rows = list(csv.reader(io.StringIO(text)))[:300]
            extracted = "\n".join(" | ".join(row[:30]) for row in rows)
            return _limit(extracted), FileExtractionStatus.TEXT_EXTRACTED, None, "已读取前 300 行。"
        if extension in {".txt", ".md"}:
            return _limit(_decode_text(data)), FileExtractionStatus.TEXT_EXTRACTED, None, ""
        if extension in {".png", ".jpg", ".jpeg", ".webp"}:
            from PIL import Image

            with Image.open(io.BytesIO(data)) as image:
                width, height = image.size
            return "", FileExtractionStatus.METADATA_ONLY, None, f"图片尺寸 {width}×{height}；本机未安装 OCR。"
        if extension in {".doc", ".xls"}:
            return "", FileExtractionStatus.PARSER_UNAVAILABLE, None, "旧版 Office 文件已留存，本机未安装转换解析器。"
    except Exception as exc:
        return "", FileExtractionStatus.FAILED, None, f"{type(exc).__name__}: {exc}"
    return "", FileExtractionStatus.PARSER_UNAVAILABLE, None, "暂无对应解析器。"


def _extract_pdf(data: bytes) -> tuple[str, FileExtractionStatus, int, str]:
    import fitz

    with fitz.open(stream=data, filetype="pdf") as document:
        page_count = document.page_count
        chunks = [document.load_page(index).get_text("text") for index in range(min(page_count, 60))]
    text = _limit("\n".join(chunks))
    if text.strip():
        note = "" if page_count <= 60 else "仅提取前 60 页文字。"
        return text, FileExtractionStatus.TEXT_EXTRACTED, page_count, note
    return "", FileExtractionStatus.METADATA_ONLY, page_count, "PDF 无可复制文字，可能是扫描件；本机未安装 OCR。"


def _extract_docx(data: bytes) -> tuple[str, FileExtractionStatus, None, str]:
    from docx import Document

    document = Document(io.BytesIO(data))
    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables[:50]:
        for row in table.rows[:300]:
            chunks.append(" | ".join(cell.text for cell in row.cells))
    for section in document.sections:
        for region in (section.header, section.footer):
            chunks.extend(
                paragraph.text
                for paragraph in region.paragraphs
                if paragraph.text.strip()
            )
    text = _limit("\n".join(chunks))
    return text, FileExtractionStatus.TEXT_EXTRACTED, None, f"已读取正文，共 {len(text)} 个字符。"


def _extract_xlsx(data: bytes) -> tuple[str, FileExtractionStatus, None, str]:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    chunks: list[str] = []
    for worksheet in workbook.worksheets[:8]:
        chunks.append(f"[{worksheet.title}]")
        for row_index, row in enumerate(worksheet.iter_rows(values_only=True)):
            if row_index >= 300:
                break
            chunks.append(" | ".join("" if value is None else str(value) for value in row[:30]))
    workbook.close()
    return _limit("\n".join(chunks)), FileExtractionStatus.TEXT_EXTRACTED, None, "每个工作表最多读取前 300 行。"


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _limit(text: str) -> str:
    return text.replace("\x00", " ")[:MAX_EXTRACTED_CHARS]


def _preview(text: str) -> str:
    compact = re.sub(r"\s+", " ", text).strip()
    return compact[:1200]


def _text_fingerprint(extension: str, text: str) -> str:
    """Deduplicate regenerated Office/PDF files whose container metadata changed."""

    normalized = _preview(text)
    if not normalized:
        return ""
    payload = f"{extension}\0{normalized}".encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _classify_evidence(
    filename: str,
    text: str,
    pending_requests: list[str],
    *,
    file_index: int,
    file_count: int,
) -> list[str]:
    filename_matches = _match_evidence_aliases(filename)
    if filename_matches:
        return filename_matches

    # Classify by the title/lead only. A contract that mentions social insurance,
    # attendance or an employee handbook does not prove those separate records exist.
    lead = "\n".join(line for line in text.splitlines()[:8] if line.strip())[:500]
    lead_matches = _match_evidence_aliases(lead)
    if lead_matches:
        return lead_matches
    if len(pending_requests) == 1:
        return [pending_requests[0]]
    if pending_requests and file_count == len(pending_requests) and file_index < len(pending_requests):
        return [pending_requests[file_index]]
    return []


def _match_evidence_aliases(haystack: str) -> list[str]:
    normalized = haystack.lower()
    matches: list[str] = []
    for evidence_name, aliases in _evidence_aliases().items():
        if any(alias.lower() in normalized for alias in (evidence_name, *aliases)):
            matches.append(evidence_name)
    return list(dict.fromkeys(matches))


def _contract_identity_quote(text: str) -> str:
    for line in text.splitlines():
        compact = line.strip()
        if compact and "劳动合同" in compact:
            return compact[:240]
    return "上传的材料被识别为劳动合同。"


def _format_extracted_fact(fact_id: str, value: object) -> str:
    labels = {
        "has_written_contract": "书面劳动合同存在",
        "contract_start_date": "合同开始日期",
        "contract_end_date": "合同结束日期",
        "contract_term_months": "合同期限",
        "probation_start_date": "试用期开始日期",
        "probation_end_date": "试用期结束日期",
        "probation_period_months": "试用期",
        "probation_monthly_salary": "试用期月工资",
        "regular_monthly_salary": "转正后月工资",
        "monthly_salary": "月工资",
    }
    label = labels.get(fact_id, fact_id)
    if fact_id.endswith("_months"):
        return f"{label}{_display_number(value)}个月"
    if "salary" in fact_id:
        return f"{label}{_display_number(value)}元"
    if fact_id == "has_written_contract":
        return label if value else "未签书面劳动合同"
    return f"{label}{value}"


def _display_number(value: object) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    return str(int(number)) if number.is_integer() else str(number)


@lru_cache(maxsize=1)
def _evidence_aliases() -> dict[str, list[str]]:
    path = Path(__file__).resolve().parent / "evidence_rules.yaml"
    with path.open("r", encoding="utf-8") as handle:
        data = yaml.safe_load(handle)
    return data.get("evidence_aliases", {})
