from io import BytesIO

import fitz
import pytest
from docx import Document
from PIL import Image

from backend.legal_domain.labor.evidence_upload import (
    CURRENT_EVIDENCE_PARSER_VERSION,
    EvidenceUploadError,
    UploadPayload,
    ingest_evidence_files,
    refresh_stored_evidence,
)
from backend.legal_domain.labor.model import get_labor_model
from backend.legal_rl.state import CaseState, FileExtractionStatus


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("劳动合同书", level=1)
    document.add_paragraph("双方签订了劳动合同，合同期限三年，试用期六个月。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _realistic_contract_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("模拟劳动合同", level=1)
    document.add_paragraph("本合同为固定期限劳动合同，合同期限自2025年09月01日起至2026年08月31日止。")
    document.add_paragraph("试用期自2025年09月01日起至2025年11月30日止，试用期共计3个月。")
    document.add_paragraph("试用期月工资为3000元（税前），转正后月基本工资为3500元（税前）。")
    document.add_paragraph("甲方依法缴纳社保，乙方遵守员工手册、考勤制度及绩效考核制度。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Payroll bank statement: monthly salary 10000")
    data = document.tobytes()
    document.close()
    return data


def _png_bytes() -> bytes:
    buffer = BytesIO()
    Image.new("RGB", (320, 180), color="white").save(buffer, format="PNG")
    return buffer.getvalue()


def test_uploads_are_stored_extracted_classified_and_merged(tmp_path):
    state = CaseState(
        dispute_type="probation_termination",
        user_narrative="公司说我试用期不合格，明天不用来了。",
    )
    get_labor_model().prepare_state(state)
    state.pending_evidence_requests = ["劳动合同", "工资流水"]

    result = ingest_evidence_files(
        state,
        [
            UploadPayload("合同扫描件.docx", _docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            UploadPayload("工资流水.pdf", _pdf_bytes(), "application/pdf"),
            UploadPayload("微信聊天记录.png", _png_bytes(), "image/png"),
        ],
        tmp_path,
    )

    assert len(result.files) == 3
    assert result.extracted_count == 2
    assert {item.extraction_status for item in result.files} == {
        FileExtractionStatus.TEXT_EXTRACTED,
        FileExtractionStatus.METADATA_ONLY,
    }
    assert {item.name for item in state.evidence} >= {"劳动合同", "工资流水", "微信聊天记录"}
    assert state.facts["has_written_contract"] is True
    assert state.facts["contract_term_months"] == 36
    assert state.facts["probation_period_months"] == 6
    assert all((tmp_path / state.case_id / item.stored_name).is_file() for item in result.files)
    assert all(item.stored_path for item in state.uploaded_files)

    duplicate = ingest_evidence_files(
        state,
        [UploadPayload("合同扫描件.docx", _docx_bytes())],
        tmp_path,
    )
    assert duplicate.duplicate_names == ["合同扫描件.docx"]
    assert len(state.uploaded_files) == 3


def test_upload_rejects_extension_spoofing(tmp_path):
    with pytest.raises(EvidenceUploadError, match="有效的 PDF"):
        ingest_evidence_files(
            CaseState(),
            [UploadPayload("伪造材料.pdf", b"not-a-pdf")],
            tmp_path,
        )


def test_contract_is_read_for_facts_without_inventing_separate_evidence_files(tmp_path):
    state = CaseState(dispute_type="probation_termination")

    result = ingest_evidence_files(
        state,
        [UploadPayload("模拟劳动合同.docx", _realistic_contract_docx_bytes())],
        tmp_path,
    )

    record = result.files[0]
    assert record.evidence_names == ["劳动合同"]
    assert {item.name for item in state.evidence} == {"劳动合同"}
    assert state.facts["has_written_contract"] is True
    assert state.facts["contract_term_months"] == 12
    assert state.facts["probation_period_months"] == 3
    assert state.facts["probation_monthly_salary"] == 3000
    assert state.facts["regular_monthly_salary"] == 3500
    assert "monthly_salary" not in state.facts
    assert "合同期限12个月" in record.extracted_facts
    assert "试用期3个月" in record.extracted_facts
    assert "转正后月工资3500元" in record.extracted_facts
    assert "材料类型：劳动合同" in result.summary_markdown()
    assert "从正文识别" in result.summary_markdown()
    document_sources = [
        item for item in state.fact_provenance
        if item.source_type == "uploaded_file" and item.source_ref == "模拟劳动合同.docx"
    ]
    assert {item.fact_id for item in document_sources} >= {
        "has_written_contract", "contract_term_months", "probation_period_months",
    }


def test_old_live_session_reindexes_stored_contract_without_reupload(tmp_path):
    state = CaseState(dispute_type="probation_termination")
    ingest_evidence_files(
        state,
        [UploadPayload("模拟劳动合同.docx", _realistic_contract_docx_bytes())],
        tmp_path,
    )
    record = state.uploaded_files[0]
    state.evidence_parser_version = 1
    record.evidence_names = ["劳动合同", "社保记录", "考勤记录", "考核记录"]
    for name in record.evidence_names[1:]:
        state.add_evidence(name, source="uploaded_file")
    state.apply_facts({"monthly_salary": 3000})
    state.add_fact_provenance(
        "monthly_salary",
        quote=record.text_preview[:120],
        extraction_method="rules",
    )

    changed = refresh_stored_evidence(state, tmp_path)

    assert changed is True
    assert state.evidence_parser_version == CURRENT_EVIDENCE_PARSER_VERSION
    assert record.evidence_names == ["劳动合同"]
    assert {item.name for item in state.evidence} == {"劳动合同"}
    assert "monthly_salary" not in state.facts
    assert state.facts["contract_term_months"] == 12
    assert "合同期限12个月" in record.extracted_facts

    assert refresh_stored_evidence(state, tmp_path) is False
